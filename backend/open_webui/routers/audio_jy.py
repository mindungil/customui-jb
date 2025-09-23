from asyncio import streams
import hashlib
import json
import logging
import os
import uuid
from functools import lru_cache
from pathlib import Path
from pydub import AudioSegment, effects
from pydub.silence import split_on_silence
from concurrent.futures import ThreadPoolExecutor
from typing import Optional
from transformers import pipeline
import torch, numpy as np
from tqdm import tqdm
import concurrent.futures
import gc
from silero_vad import load_silero_vad, read_audio, get_speech_timestamps
from df.enhance import init_df

import aiohttp
import aiofiles
import requests
import mimetypes

from fastapi import (
    Depends,
    FastAPI,
    File,
    Form,
    HTTPException,
    Request,
    UploadFile,
    status,
    APIRouter,
)
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel


from open_webui.utils.auth import get_admin_user, get_verified_user
from open_webui.config import (
    WHISPER_MODEL_AUTO_UPDATE,
    WHISPER_MODEL_DIR,
    CACHE_DIR,
    WHISPER_LANGUAGE,
)

from open_webui.constants import ERROR_MESSAGES
from open_webui.env import (
    AIOHTTP_CLIENT_SESSION_SSL,
    AIOHTTP_CLIENT_TIMEOUT,
    ENV,
    SRC_LOG_LEVELS,
    DEVICE_TYPE,
    ENABLE_FORWARD_USER_INFO_HEADERS,
)

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from transformers import pipeline as whisper_pipeline


router = APIRouter()

# Constants
MAX_FILE_SIZE_MB = 50
MAX_FILE_SIZE = MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes
AZURE_MAX_FILE_SIZE_MB = 200
AZURE_MAX_FILE_SIZE = AZURE_MAX_FILE_SIZE_MB * 1024 * 1024  # Convert MB to bytes

log = logging.getLogger(__name__)
log.setLevel(SRC_LOG_LEVELS["AUDIO"])

SPEECH_CACHE_DIR = CACHE_DIR / "audio" / "speech"
SPEECH_CACHE_DIR.mkdir(parents=True, exist_ok=True)


##########################################
#
# Utility functions
#
##########################################

from pydub import AudioSegment
from pydub.utils import mediainfo


def is_audio_conversion_required(file_path):
    """
    Check if the given audio file needs conversion to mp3.
    """
    SUPPORTED_FORMATS = {"flac", "m4a", "mp3", "mp4", "mpeg", "wav", "webm"}

    if not os.path.isfile(file_path):
        log.error(f"File not found: {file_path}")
        return False

    try:
        info = mediainfo(file_path)
        codec_name = info.get("codec_name", "").lower()
        codec_type = info.get("codec_type", "").lower()
        codec_tag_string = info.get("codec_tag_string", "").lower()

        if codec_name == "aac" and codec_type == "audio" and codec_tag_string == "mp4a":
            # File is AAC/mp4a audio, recommend mp3 conversion
            return True

        # If the codec name or file extension is in the supported formats
        if (
            codec_name in SUPPORTED_FORMATS
            or os.path.splitext(file_path)[1][1:].lower() in SUPPORTED_FORMATS
        ):
            return False  # Already supported

        return True
    except Exception as e:
        log.error(f"Error getting audio format: {e}")
        return False


def convert_audio_to_mp3(file_path):
    """Convert audio file to mp3 format."""
    try:
        output_path = os.path.splitext(file_path)[0] + ".mp3"
        audio = AudioSegment.from_file(file_path)
        audio.export(output_path, format="mp3")
        log.info(f"Converted {file_path} to {output_path}")
        return output_path
    except Exception as e:
        log.error(f"Error converting audio file: {e}")
        return None


def set_faster_whisper_model(model: str, auto_update: bool = False):
    whisper_model = None
    if model:
        from faster_whisper import WhisperModel

        faster_whisper_kwargs = {
            "model_size_or_path": model,
            "device": DEVICE_TYPE if DEVICE_TYPE and DEVICE_TYPE == "cuda" else "cpu",
            "compute_type": "int8",
            "download_root": WHISPER_MODEL_DIR,
            "local_files_only": not auto_update,
        }

        try:
            whisper_model = WhisperModel(**faster_whisper_kwargs)
        except Exception:
            log.warning(
                "WhisperModel initialization failed, attempting download with local_files_only=False"
            )
            faster_whisper_kwargs["local_files_only"] = False
            whisper_model = WhisperModel(**faster_whisper_kwargs)
    return whisper_model


##########################################
#
# Audio API
#
##########################################


class TTSConfigForm(BaseModel):
    OPENAI_API_BASE_URL: str
    OPENAI_API_KEY: str
    API_KEY: str
    ENGINE: str
    MODEL: str
    VOICE: str
    SPLIT_ON: str
    AZURE_SPEECH_REGION: str
    AZURE_SPEECH_BASE_URL: str
    AZURE_SPEECH_OUTPUT_FORMAT: str


class STTConfigForm(BaseModel):
    OPENAI_API_BASE_URL: str
    OPENAI_API_KEY: str
    ENGINE: str
    MODEL: str
    WHISPER_MODEL: str
    DEEPGRAM_API_KEY: str
    AZURE_API_KEY: str
    AZURE_REGION: str
    AZURE_LOCALES: str
    AZURE_BASE_URL: str
    AZURE_MAX_SPEAKERS: str


class AudioConfigUpdateForm(BaseModel):
    tts: TTSConfigForm
    stt: STTConfigForm


@router.get("/config")
async def get_audio_config(request: Request, user=Depends(get_admin_user)):
    return {
        "tts": {
            "OPENAI_API_BASE_URL": request.app.state.config.TTS_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.TTS_OPENAI_API_KEY,
            "API_KEY": request.app.state.config.TTS_API_KEY,
            "ENGINE": request.app.state.config.TTS_ENGINE,
            "MODEL": request.app.state.config.TTS_MODEL,
            "VOICE": request.app.state.config.TTS_VOICE,
            "SPLIT_ON": request.app.state.config.TTS_SPLIT_ON,
            "AZURE_SPEECH_REGION": request.app.state.config.TTS_AZURE_SPEECH_REGION,
            "AZURE_SPEECH_BASE_URL": request.app.state.config.TTS_AZURE_SPEECH_BASE_URL,
            "AZURE_SPEECH_OUTPUT_FORMAT": request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT,
        },
        "stt": {
            "OPENAI_API_BASE_URL": request.app.state.config.STT_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.STT_OPENAI_API_KEY,
            "ENGINE": request.app.state.config.STT_ENGINE,
            "MODEL": request.app.state.config.STT_MODEL,
            "WHISPER_MODEL": request.app.state.config.WHISPER_MODEL,
            "DEEPGRAM_API_KEY": request.app.state.config.DEEPGRAM_API_KEY,
            "AZURE_API_KEY": request.app.state.config.AUDIO_STT_AZURE_API_KEY,
            "AZURE_REGION": request.app.state.config.AUDIO_STT_AZURE_REGION,
            "AZURE_LOCALES": request.app.state.config.AUDIO_STT_AZURE_LOCALES,
            "AZURE_BASE_URL": request.app.state.config.AUDIO_STT_AZURE_BASE_URL,
            "AZURE_MAX_SPEAKERS": request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS,
        },
    }


@router.post("/config/update")
async def update_audio_config(
    request: Request, form_data: AudioConfigUpdateForm, user=Depends(get_admin_user)
):
    request.app.state.config.TTS_OPENAI_API_BASE_URL = form_data.tts.OPENAI_API_BASE_URL
    request.app.state.config.TTS_OPENAI_API_KEY = form_data.tts.OPENAI_API_KEY
    request.app.state.config.TTS_API_KEY = form_data.tts.API_KEY
    request.app.state.config.TTS_ENGINE = form_data.tts.ENGINE
    request.app.state.config.TTS_MODEL = form_data.tts.MODEL
    request.app.state.config.TTS_VOICE = form_data.tts.VOICE
    request.app.state.config.TTS_SPLIT_ON = form_data.tts.SPLIT_ON
    request.app.state.config.TTS_AZURE_SPEECH_REGION = form_data.tts.AZURE_SPEECH_REGION
    request.app.state.config.TTS_AZURE_SPEECH_BASE_URL = (
        form_data.tts.AZURE_SPEECH_BASE_URL
    )
    request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT = (
        form_data.tts.AZURE_SPEECH_OUTPUT_FORMAT
    )

    request.app.state.config.STT_OPENAI_API_BASE_URL = form_data.stt.OPENAI_API_BASE_URL
    request.app.state.config.STT_OPENAI_API_KEY = form_data.stt.OPENAI_API_KEY
    request.app.state.config.STT_ENGINE = form_data.stt.ENGINE
    request.app.state.config.STT_MODEL = form_data.stt.MODEL
    request.app.state.config.WHISPER_MODEL = form_data.stt.WHISPER_MODEL
    request.app.state.config.DEEPGRAM_API_KEY = form_data.stt.DEEPGRAM_API_KEY
    request.app.state.config.AUDIO_STT_AZURE_API_KEY = form_data.stt.AZURE_API_KEY
    request.app.state.config.AUDIO_STT_AZURE_REGION = form_data.stt.AZURE_REGION
    request.app.state.config.AUDIO_STT_AZURE_LOCALES = form_data.stt.AZURE_LOCALES
    request.app.state.config.AUDIO_STT_AZURE_BASE_URL = form_data.stt.AZURE_BASE_URL
    request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS = (
        form_data.stt.AZURE_MAX_SPEAKERS
    )

    if request.app.state.config.STT_ENGINE == "":
        request.app.state.faster_whisper_model = set_faster_whisper_model(
            form_data.stt.WHISPER_MODEL, WHISPER_MODEL_AUTO_UPDATE
        )

    return {
        "tts": {
            "OPENAI_API_BASE_URL": request.app.state.config.TTS_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.TTS_OPENAI_API_KEY,
            "API_KEY": request.app.state.config.TTS_API_KEY,
            "ENGINE": request.app.state.config.TTS_ENGINE,
            "MODEL": request.app.state.config.TTS_MODEL,
            "VOICE": request.app.state.config.TTS_VOICE,
            "SPLIT_ON": request.app.state.config.TTS_SPLIT_ON,
            "AZURE_SPEECH_REGION": request.app.state.config.TTS_AZURE_SPEECH_REGION,
            "AZURE_SPEECH_BASE_URL": request.app.state.config.TTS_AZURE_SPEECH_BASE_URL,
            "AZURE_SPEECH_OUTPUT_FORMAT": request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT,
        },
        "stt": {
            "OPENAI_API_BASE_URL": request.app.state.config.STT_OPENAI_API_BASE_URL,
            "OPENAI_API_KEY": request.app.state.config.STT_OPENAI_API_KEY,
            "ENGINE": request.app.state.config.STT_ENGINE,
            "MODEL": request.app.state.config.STT_MODEL,
            "WHISPER_MODEL": request.app.state.config.WHISPER_MODEL,
            "DEEPGRAM_API_KEY": request.app.state.config.DEEPGRAM_API_KEY,
            "AZURE_API_KEY": request.app.state.config.AUDIO_STT_AZURE_API_KEY,
            "AZURE_REGION": request.app.state.config.AUDIO_STT_AZURE_REGION,
            "AZURE_LOCALES": request.app.state.config.AUDIO_STT_AZURE_LOCALES,
            "AZURE_BASE_URL": request.app.state.config.AUDIO_STT_AZURE_BASE_URL,
            "AZURE_MAX_SPEAKERS": request.app.state.config.AUDIO_STT_AZURE_MAX_SPEAKERS,
        },
    }


def load_speech_pipeline(request):
    from transformers import pipeline
    from datasets import load_dataset

    if request.app.state.speech_synthesiser is None:
        request.app.state.speech_synthesiser = pipeline(
            "text-to-speech", "microsoft/speecht5_tts"
        )

    if request.app.state.speech_speaker_embeddings_dataset is None:
        request.app.state.speech_speaker_embeddings_dataset = load_dataset(
            "Matthijs/cmu-arctic-xvectors", split="validation"
        )


@router.post("/speech")
async def speech(request: Request, user=Depends(get_verified_user)):
    body = await request.body()
    name = hashlib.sha256(
        body
        + str(request.app.state.config.TTS_ENGINE).encode("utf-8")
        + str(request.app.state.config.TTS_MODEL).encode("utf-8")
    ).hexdigest()

    file_path = SPEECH_CACHE_DIR.joinpath(f"{name}.mp3")
    file_body_path = SPEECH_CACHE_DIR.joinpath(f"{name}.json")

    # Check if the file already exists in the cache
    if file_path.is_file():
        return FileResponse(file_path)

    payload = None
    try:
        payload = json.loads(body.decode("utf-8"))
    except Exception as e:
        log.exception(e)
        raise HTTPException(status_code=400, detail="Invalid JSON payload")

    if request.app.state.config.TTS_ENGINE == "openai":
        payload["model"] = request.app.state.config.TTS_MODEL

        try:
            timeout = aiohttp.ClientTimeout(total=AIOHTTP_CLIENT_TIMEOUT)
            async with aiohttp.ClientSession(
                timeout=timeout, trust_env=True
            ) as session:
                async with session.post(
                    url=f"{request.app.state.config.TTS_OPENAI_API_BASE_URL}/audio/speech",
                    json=payload,
                    headers={
                        "Content-Type": "application/json",
                        "Authorization": f"Bearer {request.app.state.config.TTS_OPENAI_API_KEY}",
                        **(
                            {
                                "X-OpenWebUI-User-Name": user.name,
                                "X-OpenWebUI-User-Id": user.id,
                                "X-OpenWebUI-User-Email": user.email,
                                "X-OpenWebUI-User-Role": user.role,
                            }
                            if ENABLE_FORWARD_USER_INFO_HEADERS
                            else {}
                        ),
                    },
                    ssl=AIOHTTP_CLIENT_SESSION_SSL,
                ) as r:
                    r.raise_for_status()

                    async with aiofiles.open(file_path, "wb") as f:
                        await f.write(await r.read())

                    async with aiofiles.open(file_body_path, "w") as f:
                        await f.write(json.dumps(payload))

            return FileResponse(file_path)

        except Exception as e:
            log.exception(e)
            detail = None

            try:
                if r.status != 200:
                    res = await r.json()

                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
            except Exception:
                detail = f"External: {e}"

            raise HTTPException(
                status_code=getattr(r, "status", 500) if r else 500,
                detail=detail if detail else "Open WebUI: Server Connection Error",
            )

    elif request.app.state.config.TTS_ENGINE == "elevenlabs":
        voice_id = payload.get("voice", "")

        if voice_id not in get_available_voices(request):
            raise HTTPException(
                status_code=400,
                detail="Invalid voice id",
            )

        try:
            timeout = aiohttp.ClientTimeout(total=AIOHTTP_CLIENT_TIMEOUT)
            async with aiohttp.ClientSession(
                timeout=timeout, trust_env=True
            ) as session:
                async with session.post(
                    f"https://api.elevenlabs.io/v1/text-to-speech/{voice_id}",
                    json={
                        "text": payload["input"],
                        "model_id": request.app.state.config.TTS_MODEL,
                        "voice_settings": {"stability": 0.5, "similarity_boost": 0.5},
                    },
                    headers={
                        "Accept": "audio/mpeg",
                        "Content-Type": "application/json",
                        "xi-api-key": request.app.state.config.TTS_API_KEY,
                    },
                    ssl=AIOHTTP_CLIENT_SESSION_SSL,
                ) as r:
                    r.raise_for_status()

                    async with aiofiles.open(file_path, "wb") as f:
                        await f.write(await r.read())

                    async with aiofiles.open(file_body_path, "w") as f:
                        await f.write(json.dumps(payload))

            return FileResponse(file_path)

        except Exception as e:
            log.exception(e)
            detail = None

            try:
                if r.status != 200:
                    res = await r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
            except Exception:
                detail = f"External: {e}"

            raise HTTPException(
                status_code=getattr(r, "status", 500) if r else 500,
                detail=detail if detail else "Open WebUI: Server Connection Error",
            )

    elif request.app.state.config.TTS_ENGINE == "azure":
        try:
            payload = json.loads(body.decode("utf-8"))
        except Exception as e:
            log.exception(e)
            raise HTTPException(status_code=400, detail="Invalid JSON payload")

        region = request.app.state.config.TTS_AZURE_SPEECH_REGION or "eastus"
        base_url = request.app.state.config.TTS_AZURE_SPEECH_BASE_URL
        language = request.app.state.config.TTS_VOICE
        locale = "-".join(request.app.state.config.TTS_VOICE.split("-")[:1])
        output_format = request.app.state.config.TTS_AZURE_SPEECH_OUTPUT_FORMAT

        try:
            data = f"""<speak version="1.0" xmlns="http://www.w3.org/2001/10/synthesis" xml:lang="{locale}">
                <voice name="{language}">{payload["input"]}</voice>
            </speak>"""
            timeout = aiohttp.ClientTimeout(total=AIOHTTP_CLIENT_TIMEOUT)
            async with aiohttp.ClientSession(
                timeout=timeout, trust_env=True
            ) as session:
                async with session.post(
                    (base_url or f"https://{region}.tts.speech.microsoft.com")
                    + "/cognitiveservices/v1",
                    headers={
                        "Ocp-Apim-Subscription-Key": request.app.state.config.TTS_API_KEY,
                        "Content-Type": "application/ssml+xml",
                        "X-Microsoft-OutputFormat": output_format,
                    },
                    data=data,
                    ssl=AIOHTTP_CLIENT_SESSION_SSL,
                ) as r:
                    r.raise_for_status()

                    async with aiofiles.open(file_path, "wb") as f:
                        await f.write(await r.read())

                    async with aiofiles.open(file_body_path, "w") as f:
                        await f.write(json.dumps(payload))

                    return FileResponse(file_path)

        except Exception as e:
            log.exception(e)
            detail = None

            try:
                if r.status != 200:
                    res = await r.json()
                    if "error" in res:
                        detail = f"External: {res['error'].get('message', '')}"
            except Exception:
                detail = f"External: {e}"

            raise HTTPException(
                status_code=getattr(r, "status", 500) if r else 500,
                detail=detail if detail else "Open WebUI: Server Connection Error",
            )

    elif request.app.state.config.TTS_ENGINE == "transformers":
        payload = None
        try:
            payload = json.loads(body.decode("utf-8"))
        except Exception as e:
            log.exception(e)
            raise HTTPException(status_code=400, detail="Invalid JSON payload")

        import torch
        import soundfile as sf

        load_speech_pipeline(request)

        embeddings_dataset = request.app.state.speech_speaker_embeddings_dataset

        speaker_index = 6799
        try:
            speaker_index = embeddings_dataset["filename"].index(
                request.app.state.config.TTS_MODEL
            )
        except Exception:
            pass

        speaker_embedding = torch.tensor(
            embeddings_dataset[speaker_index]["xvector"]
        ).unsqueeze(0)

        speech = request.app.state.speech_synthesiser(
            payload["input"],
            forward_params={"speaker_embeddings": speaker_embedding},
        )

        sf.write(file_path, speech["audio"], samplerate=speech["sampling_rate"])

        async with aiofiles.open(file_body_path, "w") as f:
            await f.write(json.dumps(payload))

        return FileResponse(file_path)

def filter_repetitive_text(text, min_repeat=3, max_repeat_ratio=0.7):
    """
    반복되는 텍스트를 필터링하는 함수
    - min_repeat: 연속 반복 최소 횟수
    - max_repeat_ratio: 전체 텍스트에서 반복 부분이 차지하는 최대 비율
    """
    if not text or len(text.strip()) == 0:
        return ""

    words = text.split()
    if len(words) < min_repeat:
        return text

    # 1. 연속된 동일 단어 제거
    filtered_words = []
    prev_word = None
    repeat_count = 0

    for word in words:
        if word == prev_word:
            repeat_count += 1
            if repeat_count < min_repeat:  # 3번까지는 허용
                filtered_words.append(word)
        else:
            filtered_words.append(word)
            repeat_count = 0
        prev_word = word

    # 2. 전체 텍스트가 반복으로만 이루어져 있는지 확인
    filtered_text = " ".join(filtered_words)
    unique_words = set(filtered_words)

    # 고유 단어가 너무 적으면 (반복이 심하면) 대표 문구만 남김
    if len(unique_words) <= 3 and len(filtered_words) > 10:
        # 가장 흔한 단어 조합 찾기
        from collections import Counter
        word_pairs = [f"{filtered_words[i]} {filtered_words[i+1]}"
                     for i in range(len(filtered_words)-1)]
        most_common = Counter(word_pairs).most_common(1)

        if most_common and most_common[0][1] > len(filtered_words) * 0.3:
            return most_common[0][0]  # 가장 흔한 조합만 반환

    return filtered_text


def control_RMS(input_path):
    try:
        audio = AudioSegment.from_file(input_path)
        
        avg_rms = audio.dBFS
        if avg_rms < -35:   # 전체가 너무 작다?
            # 10 dB 범위 안에서만 한 번 전체 부스트
            gain = min(-25 - avg_rms, 10)
            audio = audio.apply_gain(gain)
        
        log.info(audio.dBFS)
        audio.export(input_path, format="wav")
        return input_path
    
    except Exception as e:
        log.error(f"WAV 변환 실패: {str(e)}")
        raise



def convert_to_wav(input_path):
    try:
        audio = AudioSegment.from_file(input_path)
        log.info(audio.dBFS)

                
        # 노이즈 제거 및 오디오 품질 개선
        audio = effects.normalize(audio, headroom=3.0)
        log.info(audio.dBFS)
        # 볼륨 부스트 (약한 소리도 잘 인식하도록)
        audio = audio + 2.5  # dB 부스트

        TARGET_RMS = -20.0
        PEAK_LIMIT = -1.0
        # ── 1) 피크를 -1 dBFS 이하로 한 번 정렬 ──────────────────
        audio = effects.normalize(audio, headroom=abs(PEAK_LIMIT))

        # ── 2) 현재 RMS를 계산해 목표 RMS(-20 dBFS)까지 게인 적용 ──
        current_rms  = audio.dBFS
        gain_needed  = TARGET_RMS - current_rms   # dB 단위
        audio        = audio.apply_gain(gain_needed)

        # 게인 후 피크가 다시 튈 수 있으므로 한 번 더 체크
        if audio.max_dBFS > PEAK_LIMIT:
            audio = effects.normalize(audio, headroom=abs(PEAK_LIMIT))

        
        # 샘플레이트 조정 (16kHz가 Whisper에 최적)
        audio = audio.set_frame_rate(16000)
        log.info(audio.dBFS)
        # 모노로 변환 (스테레오는 Whisper에서 불필요)
        audio = audio.set_channels(1)

        output_path = os.path.splitext(input_path)[0] + ".wav"
        audio.export(output_path, format="wav")

        log.info(f"WAV 변환 완료: {output_path}")
        return output_path
    except Exception as e:
        log.error(f"WAV 변환 실패: {str(e)}")
        raise

# pip install silero-vad    ← 한 번만 설치
def split_audio_with_timestamps_vad(
        audio_path: str,
        max_chunk_ms: int = 30000,       # Whisper 입력 30 s
        min_chunk_ms: int = 100,        # 0.1s 미만 버림
        padding_ms: int = 200):          # 경계 패딩(±0.2 s)
    """
    Silero VAD 기반 오디오 분할:
      • 음성 구간을 NN-VAD로 탐지
      • 너무 긴 구간(>20 s)은 20 s 단위 슬라이스
      • 각 청크를 WAV로 저장 + 절대 타임스탬프 반환
    """
    try:
        # 0) 오디오 → PyTorch Tensor (16 kHz mono PCM  in float32, -1..1)
        full_audio = AudioSegment.from_wav(audio_path)
        samples = np.array(full_audio.get_array_of_samples()).astype(np.float32) / 32768
        wav_tensor = torch.from_numpy(samples).unsqueeze(0)  # (1, T)

        # 1) silero-vad 모델 로드 & 추론
        model = load_silero_vad()                 # 첫 호출 시 다운로드
        speech_segments = get_speech_timestamps(
            wav_tensor, model,
            threshold = 0.16,
            speech_pad_ms=padding_ms,
            min_speech_duration_ms= 250,
            min_silence_duration_ms=500)      # ← 인자로 받은 padding_ms 활용
        

        log.info(speech_segments)
        speech_segments = [[seg["start"] / 16, seg["end"] / 16] for seg in speech_segments]
        
        if not speech_segments:
            log.warning("음성 구간이 탐지되지 않았습니다.")
            return []
        log.info(speech_segments)
        
        merged = []
        for s, e in merged:
            if merged and s - merged[-1][1] < 120:           # gap < 0.25 s
                merged[-1][1] = e
            else:
                merged.append([s, e])
        
        chunk_info = []

        for start_ms, end_ms in speech_segments:
            # 실제 음성 구간 추출
            chunk = full_audio[start_ms:end_ms]

            chunk_info.append({
                'chunk': chunk,
                'start_ms': start_ms,
                'end_ms': end_ms,
                'duration_ms': len(chunk)
            })

        # 너무 긴 청크는 나누기 (시간 정보 유지)
        final_chunk_info = []
        for info in chunk_info:
            if info['duration_ms'] > max_chunk_ms:
                chunk = info['chunk']
                start_time = info['start_ms']

                # 최대 청크 길이로 나누기
                for i in range(0, len(chunk), max_chunk_ms):
                    sliced = chunk[i:i + max_chunk_ms]
                    if len(sliced) >= min_chunk_ms:
                        slice_start = start_time + i
                        slice_end = slice_start + len(sliced)

                        final_chunk_info.append({
                            'chunk': sliced,
                            'start_ms': slice_start,
                            'end_ms': slice_end,
                            'duration_ms': len(sliced)
                        })
            else:
                final_chunk_info.append(info)

        # 분할된 청크 WAV로 저장
        chunk_paths_with_time = []
        for i, info in enumerate(final_chunk_info):
            chunk_path = f"{os.path.splitext(audio_path)[0]}_chunk_{i}.wav"
            info['chunk'].export(chunk_path, format="wav")

            chunk_paths_with_time.append({
                'path': chunk_path,
                'start_ms': info['start_ms'],
                'end_ms': info['end_ms'],
                'duration_ms': info['duration_ms']
            })

        log.info(chunk_paths_with_time)
        log.info(f"총 {len(chunk_paths_with_time)}개의 청크로 분할 완료 (실제 타임스탬프 추적)")
        return chunk_paths_with_time

    except Exception as e:
        log.error(f"오디오 분할 실패: {str(e)}")
        raise


# 한 청크 Whisper STT 수행 (타임스탬프 정보 포함)
def process_chunk_with_timestamp(chunk_info, pipe):
    """
    개별 청크 처리 함수 - 실제 타임스탬프 반환
    """
    try:
        chunk_path = chunk_info['path']
        start_ms = chunk_info['start_ms']

        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gc.collect()

        # 청크 처리 전에 오디오 정규화
        audio = AudioSegment.from_wav(chunk_path)

        # 오디오가 너무 조용하거나 단조로운지 확인
        if audio.dBFS < -60:
            log.info(f"청크 {chunk_path}가 너무 조용함, 건너뜀")
            return None

        temp_path = f"{chunk_path}_normalized.wav"
        audio.export(temp_path, format="wav")

        # Whisper 처리 시 타임스탬프 반환 옵션 활성화
        # Whisper 파라미터 최적화 (반복 방지와 안정성 강화)
        result = pipe(
            temp_path,
            stride_length_s=7,           # 오버랩 구간
            batch_size=8,               # 배치 크기 감소로 안정성 확보
            return_timestamps=True,      # 타임스탬프 반환
            generate_kwargs={
                "max_new_tokens": 256,   # 토큰 수 제한으로 환각 방지
                "temperature": 0.0,      # 결정적 생성
                "do_sample": False,      # 샘플링 비활성화
                "repetition_penalty": 1.2,  # 반복 방지
                "no_repeat_ngram_size": 3,  # 3-gram 반복 방지
                "num_beams": 1,            # 빔 서치 비활성화로 빠른 처리
                "language": "ko",        # korean으로 설정
                "task": "transcribe",    # 텍스트로 그대로 변환
            }
        )

        text = result.get('text', '').strip()

        # 반복 문구 필터링 적용
        '''
        filtered_text = filter_repetitive_text(text)

        # 추가 품질 검사
        if len(filtered_text.split()) < 2:
            log.info(f"청크 {chunk_path}에서 의미있는 텍스트 없음")
            return None
        '''

        # 임시 파일 삭제
        try:
            os.remove(temp_path)
        except:
            pass

        # 실제 오디오 파일에서의 절대 시간 계산
        return {
            'text': text,
            'start_ms': start_ms,
            'end_ms': chunk_info['end_ms'],
            'chunk_timestamps': result.get('chunks', [])  # Whisper 내부 타임스탬프
        }

    except Exception as e:
        log.error(f"청크 처리 실패 ({chunk_info['path']}): {str(e)}")
        return None

def convert_seconds_to_hms(seconds):
    """Convert seconds to HH:MM:SS format"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    seconds = int(seconds % 60)
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"

# 전체 진행 함수
def transcribe_long_audio(request: Request, file_path, model_name='openai/whisper-large-v3'):
    try:
        # GPU 리소스 정리
        if torch.cuda.is_available():
            torch.cuda.empty_cache()
            gc.collect()
            torch.backends.cudnn.benchmark = True

        # 오디오 WAV로 변환
        wav_path = convert_to_wav(file_path)

        # 전체 오디오 길이 계산
        full_audio = AudioSegment.from_wav(wav_path)
        total_duration_ms = len(full_audio)

        # 디바이스 설정 (GPU or CPU)
        # GPU 변경 후 사용 가능한 GPU로 사용하는 자동화가 필요합니다.
        device = 1 if torch.cuda.is_available() else -1
        log.info(f"Whisper 실행 중 (Device: {'GPU' if device == 1 else 'CPU'})")

        # Whisper 파이프라인 로딩
        log.info("Whisper 파이프라인 로딩 중...")
        whisper = whisper_pipeline(
            "automatic-speech-recognition",
            model=model_name,
            device=device
        )

        # 오디오 청크로 분할 (타임스탬프 정보 포함)
        log.info("오디오 파일을 청크로 분할 중...")
        chunk_infos = split_audio_with_timestamps_vad(wav_path)

        # 청크별 STT 처리
        log.info("STT 처리 중...")
        results = []
        segments = []
        batch_size = 3

        for i in range(0, len(chunk_infos), batch_size):
            batch_chunk_infos = chunk_infos[i:i + batch_size]
            batch_results = []

            with concurrent.futures.ThreadPoolExecutor(max_workers=2) as executor:
                future_to_chunk = {
                    executor.submit(process_chunk_with_timestamp, chunk_info, whisper): chunk_info
                    for chunk_info in batch_chunk_infos
                }

                for future in tqdm(concurrent.futures.as_completed(future_to_chunk), total=len(batch_chunk_infos)):
                    chunk_info = future_to_chunk[future]
                    try:
                        result = future.result()

                        if result is not None:
                            batch_results.append(result)
                            results.append(result['text'])

                            # 실제 타임스탬프로 세그먼트 추가
                            segments.append({
                                "start": convert_seconds_to_hms(result['start_ms'] / 1000.0),
                                "end": convert_seconds_to_hms(result['end_ms'] / 1000.0),
                                "text": result['text'].strip(),
                                "start_seconds": result['start_ms'] / 1000.0,
                                "end_seconds": result['end_ms'] / 1000.0
                            })

                    except Exception as e:
                        log.error(f"청크 처리 실패: {str(e)}")

            if torch.cuda.is_available():
                torch.cuda.empty_cache()
                gc.collect()

        # 세그먼트를 시간순으로 정렬
        segments.sort(key=lambda x: x['start_seconds'])

        # 전체 텍스트 통합 및 최종 정리
        plain_text = " ".join([seg['text'] for seg in segments])
        
        # docx 문서 생성
        doc = Document()
        doc.add_heading('회의록', 0)

        # 스타일 설정
        style = doc.styles['Normal']
        style.font.name = '맑은 고딕'
        style.font.size = Pt(11)

        # 시간순으로 정렬된 세그먼트로 스크립트 생성
        for segment in segments:
            p = doc.add_paragraph()
            p.add_run(f"[{segment['start']} - {segment['end']}] ").bold = True
            p.add_run(segment['text'])

        # 임시 청크 파일 삭제
        for chunk_info in chunk_infos:
            try:
                os.remove(chunk_info['path'])
            except:
                pass

        log.info("STT 처리 완료")
        return {
            'plain_text': plain_text,
            'docx_document': doc,
            'segments': segments
        }

    except Exception as e:
        log.error(f"STT 처리 실패: {str(e)}")
        raise

def transcription_handler(request, file_path, metadata):
    filename = os.path.basename(file_path)
    file_dir = os.path.dirname(file_path)
    id = filename.split(".")[0]

    metadata = metadata or {}

    if request.app.state.config.STT_ENGINE == "":
        if request.app.state.faster_whisper_model is None:
            request.app.state.faster_whisper_model = set_faster_whisper_model(
                request.app.state.config.WHISPER_MODEL
            )

        model = request.app.state.faster_whisper_model
        segments, info = model.transcribe(
            file_path,
            beam_size=5,
            vad_filter=request.app.state.config.WHISPER_VAD_FILTER,
            language=metadata.get("language") or WHISPER_LANGUAGE,
        )
        log.info(
            "Detected language '%s' with probability %f"
            % (info.language, info.language_probability)
        )

        # 시간대별 세그먼트 정보를 포함한 데이터 생성
        segments_list = []
        for segment in segments:
            segments_list.append({
                "start": round(segment.start, 2),
                "end": round(segment.end, 2),
                "text": segment.text.strip()
            })

        data = {
            "text": "".join([segment["text"] for segment in segments_list]),
            "segments": segments_list
        }

        # save the transcript to a json file
        transcript_file = f"{file_dir}/{id}.json"
        with open(transcript_file, "w") as f:
            json.dump(data, f)

        log.debug(data)
        return data


def transcribe(request: Request, file_path: str, metadata: Optional[dict] = None, filedata: list = None):
    log.info(f"transcribe: {file_path} {metadata}")
    log.info(f"filedata: {filedata}")

    result = transcribe_long_audio(request, file_path)
    plain_text = result['plain_text']
    docx_doc = result['docx_document']
    segments = result['segments']

    try:
        # 1. 기존 txt 파일 저장 (변경 없음)
        save_path = os.path.join(os.path.dirname(file_path), f"{os.path.splitext(filedata[1])[0]}.txt")
        with open(save_path, "w", encoding="utf-8") as f:
            f.write(plain_text)
        log.info(f"Transcript saved to: {save_path}")

        from open_webui.models.files import FileForm, Files
        from open_webui.storage.provider import Storage
        import uuid

        # 2. txt 파일을 영구 저장소에 업로드 (변경 없음)
        txt_id = f"{filedata[0]}txt"
        txt_name = f"{os.path.splitext(filedata[1])[0]}txt"
        txt_filename = f"{filedata[2]}txt"
        with open(save_path, "rb") as f:
            txt_bytes, txt_storage_path = Storage.upload_file(
                f, txt_filename,
                tags = {
                    **filedata[3],
                    "OpenWebUI-File-Id": txt_id,
                    "OpenWebUI-Transcript-Of": f"{filedata[0]}"
                }
            )

        user_id = filedata[3]["OpenWebUI-User-Id"]
        # 3. Files 테이블에 txt 파일 row 생성 (변경 없음)
        Files.insert_new_file(
            user_id,
            FileForm(
                id = txt_id,
                filename = txt_name,
                path = txt_storage_path,
                meta = {
                    "name": txt_name,
                    "content_type": "text/plain",
                    "size": len(txt_bytes),
                    "data": {"transcript_of": filedata[0]}
                },
            )
        )

        # 4. docx 파일 저장 및 업로드 (새로 추가)
        docx_path = os.path.join(os.path.dirname(file_path), f"{os.path.splitext(filedata[1])[0]}.docx")
        docx_doc.save(docx_path)
        log.info(f"Detailed transcript saved to: {docx_path}")

        # 5. docx 파일을 영구 저장소에 업로드
        docx_id = f"{filedata[0]}docx"
        docx_name = f"{os.path.splitext(filedata[1])[0]}.docx"
        docx_filename = f"{filedata[2]}.docx"
        with open(docx_path, "rb") as f:
            docx_bytes, docx_storage_path = Storage.upload_file(
                f, docx_filename,
                tags = {
                    **filedata[3],
                    "OpenWebUI-File-Id": docx_id,
                    "OpenWebUI-Transcript-Of": f"{filedata[0]}",
                    "OpenWebUI-Transcript-Type": "detailed"
                }
            )

        # 6. Files 테이블에 docx 파일 row 생성
        Files.insert_new_file(
            user_id,
            FileForm(
                id = docx_id,
                filename = docx_name,
                path = docx_storage_path,
                meta = {
                    "name": docx_name,
                    "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "size": len(docx_bytes),
                    "data": {
                        "transcript_of": filedata[0],
                        "transcript_type": "detailed",
                        "segments": segments
                    }
                },
            )
        )

        return {
            "text": plain_text,
            "transcript_file_id": txt_id,
            "detailed_transcript_file_id": docx_id
        }

    except Exception as e:
        log.exception(f"Failed to save or upload transcript: {e}")
        return {
            "text": plain_text,
            "error": f"Transcript saved locally, but upload failed: {str(e)}",
        }



@router.post("/transcriptions")
def transcription(
    request: Request,
    file: UploadFile = File(...),
    language: Optional[str] = Form(None),
    user=Depends(get_verified_user),
):
    log.info(f"file.content_type: {file.content_type}")

    SUPPORTED_CONTENT_TYPES = {"video/webm"}  # Extend if you add more video types!
    if not (
        file.content_type.startswith("audio/")
        or file.content_type in SUPPORTED_CONTENT_TYPES
    ):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.FILE_NOT_SUPPORTED,
        )

    try:
        ext = file.filename.split(".")[-1]
        id = uuid.uuid4()

        filename = f"{id}.{ext}"
        contents = file.file.read()

        file_dir = f"{CACHE_DIR}/audio/transcriptions"
        os.makedirs(file_dir, exist_ok=True)
        file_path = f"{file_dir}/{filename}"

        with open(file_path, "wb") as f:
            f.write(contents)

        try:
            metadata = None

            if language:
                metadata = {"language": language}

            result = transcribe(request, file_path, metadata)

            return {
                **result,
                "filename": os.path.basename(file_path),
            }

        except Exception as e:
            log.exception(e)

            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=ERROR_MESSAGES.DEFAULT(e),
            )

    except Exception as e:
        log.exception(e)

        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=ERROR_MESSAGES.DEFAULT(e),
        )


def get_available_models(request: Request) -> list[dict]:
    available_models = []
    if request.app.state.config.TTS_ENGINE == "openai":
        # Use custom endpoint if not using the official OpenAI API URL
        if not request.app.state.config.TTS_OPENAI_API_BASE_URL.startswith(
            "https://api.openai.com"
        ):
            try:
                response = requests.get(
                    f"{request.app.state.config.TTS_OPENAI_API_BASE_URL}/audio/models"
                )
                response.raise_for_status()
                data = response.json()
                available_models = data.get("models", [])
            except Exception as e:
                log.error(f"Error fetching models from custom endpoint: {str(e)}")
                available_models = [{"id": "tts-1"}, {"id": "tts-1-hd"}]
        else:
            available_models = [{"id": "tts-1"}, {"id": "tts-1-hd"}]
    elif request.app.state.config.TTS_ENGINE == "elevenlabs":
        try:
            response = requests.get(
                "https://api.elevenlabs.io/v1/models",
                headers={
                    "xi-api-key": request.app.state.config.TTS_API_KEY,
                    "Content-Type": "application/json",
                },
                timeout=5,
            )
            response.raise_for_status()
            models = response.json()

            available_models = [
                {"name": model["name"], "id": model["model_id"]} for model in models
            ]
        except requests.RequestException as e:
            log.error(f"Error fetching voices: {str(e)}")
    return available_models


@router.get("/models")
async def get_models(request: Request, user=Depends(get_verified_user)):
    return {"models": get_available_models(request)}


def get_available_voices(request) -> dict:
    """Returns {voice_id: voice_name} dict"""
    available_voices = {}
    if request.app.state.config.TTS_ENGINE == "openai":
        # Use custom endpoint if not using the official OpenAI API URL
        if not request.app.state.config.TTS_OPENAI_API_BASE_URL.startswith(
            "https://api.openai.com"
        ):
            try:
                response = requests.get(
                    f"{request.app.state.config.TTS_OPENAI_API_BASE_URL}/audio/voices"
                )
                response.raise_for_status()
                data = response.json()
                voices_list = data.get("voices", [])
                available_voices = {voice["id"]: voice["name"] for voice in voices_list}
            except Exception as e:
                log.error(f"Error fetching voices from custom endpoint: {str(e)}")
                available_voices = {
                    "alloy": "alloy",
                    "echo": "echo",
                    "fable": "fable",
                    "onyx": "onyx",
                    "nova": "nova",
                    "shimmer": "shimmer",
                }
        else:
            available_voices = {
                "alloy": "alloy",
                "echo": "echo",
                "fable": "fable",
                "onyx": "onyx",
                "nova": "nova",
                "shimmer": "shimmer",
            }
    elif request.app.state.config.TTS_ENGINE == "elevenlabs":
        try:
            available_voices = get_elevenlabs_voices(
                api_key=request.app.state.config.TTS_API_KEY
            )
        except Exception:
            # Avoided @lru_cache with exception
            pass
    elif request.app.state.config.TTS_ENGINE == "azure":
        try:
            region = request.app.state.config.TTS_AZURE_SPEECH_REGION
            base_url = request.app.state.config.TTS_AZURE_SPEECH_BASE_URL
            url = (
                base_url or f"https://{region}.tts.speech.microsoft.com"
            ) + "/cognitiveservices/voices/list"
            headers = {
                "Ocp-Apim-Subscription-Key": request.app.state.config.TTS_API_KEY
            }

            response = requests.get(url, headers=headers)
            response.raise_for_status()
            voices = response.json()

            for voice in voices:
                available_voices[voice["ShortName"]] = (
                    f"{voice['DisplayName']} ({voice['ShortName']})"
                )
        except requests.RequestException as e:
            log.error(f"Error fetching voices: {str(e)}")

    return available_voices


@lru_cache
def get_elevenlabs_voices(api_key: str) -> dict:
    """
    Note, set the following in your .env file to use Elevenlabs:
    AUDIO_TTS_ENGINE=elevenlabs
    AUDIO_TTS_API_KEY=sk_...  # Your Elevenlabs API key
    AUDIO_TTS_VOICE=EXAVITQu4vr4xnSDxMaL  # From https://api.elevenlabs.io/v1/voices
    AUDIO_TTS_MODEL=eleven_multilingual_v2
    """

    try:
        # TODO: Add retries
        response = requests.get(
            "https://api.elevenlabs.io/v1/voices",
            headers={
                "xi-api-key": api_key,
                "Content-Type": "application/json",
            },
        )
        response.raise_for_status()
        voices_data = response.json()

        voices = {}
        for voice in voices_data.get("voices", []):
            voices[voice["voice_id"]] = voice["name"]
    except requests.RequestException as e:
        # Avoid @lru_cache with exception
        log.error(f"Error fetching voices: {str(e)}")
        raise RuntimeError(f"Error fetching voices: {str(e)}")

    return voices


@router.get("/voices")
async def get_voices(request: Request, user=Depends(get_verified_user)):
    return {
        "voices": [
            {"id": k, "name": v} for k, v in get_available_voices(request).items()
        ]
    }
